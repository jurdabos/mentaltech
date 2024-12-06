from docx import Document
import featuretools as ft
import fitz
from hdbscan import HDBSCAN
import numpy as np
import pandas as pd
from scipy import stats
from scipy.spatial.distance import cdist
from sklearn.feature_selection import VarianceThreshold
from sklearn.inspection import permutation_importance
from sklearn.metrics import pairwise_distances, silhouette_score
from sklearn.neighbors import KNeighborsClassifier
from sklearn.preprocessing import RobustScaler
import tabulate


def automated_feature_engineering_and_clustering(
        df,
        entityset_name,
        max_depth=1,
        robust_scaling=True,
        variance_threshold=0.05,
        corr_threshold=0.8,
        min_cluster_size=10,
        min_samples=5
):
    """
    Performs DFS, scaling, feature selection, HDBSCAN clustering, and feature importance calculation.
    This is a skeleton for future workflows involving higher levels of abstraction.
    Parameters:
    - df: pd.DataFrame - Input data
    - entityset_name: str - Name for the EntitySet
    - max_depth: int - Maximum depth for DFS
    - robust_scaling: bool - Whether to apply robust scaling
    - variance_threshold: float - Threshold for variance filtering
    - corr_threshold: float - Threshold for correlation deselection
    - min_cluster_size: int - HDBSCAN minimum cluster size
    - min_samples: int - HDBSCAN minimum samples
    Returns:
    - final_features: pd.DataFrame - Feature-engineered and selected dataset
    - clusters: np.array - Cluster labels
    - feature_importances: pd.DataFrame - Feature importance ranking
    """
    # Fixed primitives
    trans_primitives = [
        "add_numeric", "subtract_numeric", "divide_numeric", "multiply_numeric"
    ]
    agg_primitives = [
        "mean", "std", "min", "max", "count"
    ]
    # Step 1: Add EntitySet and DFS
    df['index_col'] = range(len(df))
    es = ft.EntitySet(id=entityset_name)
    es = es.add_dataframe(
        dataframe_name='TableEntity',
        dataframe=df,
        index='index_col'
    )
    features, feature_names = ft.dfs(
        entityset=es,
        target_dataframe_name='TableEntity',
        max_depth=max_depth,
        trans_primitives=trans_primitives,
        agg_primitives=agg_primitives
    )
    features.reset_index(drop=True, inplace=True)
    # Step 2: Handle NaN/Inf Values and Remove Index Column
    features.replace([np.inf, -np.inf], np.nan, inplace=True)
    features.dropna(axis=1, inplace=True)
    if 'index_col' in features.columns:
        features.drop(columns=['index_col'], inplace=True)
    # Step 3: Robust Scaling
    if robust_scaling:
        scaler = RobustScaler()
        scaled_features = pd.DataFrame(scaler.fit_transform(features), columns=features.columns)
    else:
        scaled_features = features
    # Step 4: Variance Thresholding
    vt = VarianceThreshold(threshold=variance_threshold)
    vt_features = pd.DataFrame(vt.fit_transform(scaled_features), columns=scaled_features.columns[vt.get_support()])
    # Step 5: Correlation-Based Feature Selection
    corr_matrix = vt_features.corr(method='spearman')
    upper_triangle = corr_matrix.where(np.triu(np.ones(corr_matrix.shape), k=1).astype(bool))
    to_drop = [col for col in upper_triangle.columns if any(upper_triangle[col] > corr_threshold)]
    selected_features = vt_features.drop(columns=to_drop)
    # Step 6: HDBSCAN Clustering
    clusterer = HDBSCAN(min_cluster_size=min_cluster_size, min_samples=min_samples)
    clusterer.fit(selected_features)
    selected_features['Cluster'] = clusterer.labels_
    selected_features['Probability'] = clusterer.probabilities_
    # Step 7: Feature Importance Using Clusters as Target
    X = selected_features.drop(columns=['Cluster', 'Probability'])
    y = selected_features['Cluster']
    model = KNeighborsClassifier()
    model.fit(X, y)
    feat_imp = permutation_importance(model, X, y, n_repeats=10, scoring='accuracy', random_state=42)
    feature_importances = pd.DataFrame({
        'features': X.columns,
        'importances_mean': feat_imp.importances_mean,
        'importances_std': feat_imp.importances_std
    }).sort_values(by='importances_mean', ascending=False)
    # Step 8: Select Top Features
    top_features = feature_importances.head(8)['features']
    final_features = selected_features[top_features]
    return final_features, clusterer.labels_, feature_importances


def calculate_clustering_metrics(name, labels, data, cluster_centers=None, model=None):
    """
    Calculates clustering QA metrics: Noise Percentage, Silhouette Score, weighted WSS, and BIC.
    Parameters:
    - name: str, clustering experiment identifier
    - labels: array-like, cluster labels (noise points should be -1 for DBSCAN)
    - data: array-like, original dataset
    - cluster_centers: array-like, optional cluster centers (needed for WSS)
    - model: clustering model object (used for BIC if supported)
    Returns:
    - dict with clustering metrics
    """
    non_noise_indices = labels != -1
    clustered_data = data[non_noise_indices].to_numpy()
    clustered_labels = labels[non_noise_indices]
    # Calculating noise percentage
    noise_percentage = 1 - np.sum(non_noise_indices, axis=0) / len(labels)
    # Calculating silhouette score
    if len(np.unique(clustered_labels)) > 1:
        silhouette = silhouette_score(clustered_data, clustered_labels)
    else:
        silhouette = np.nan
    # Calculating WSS
    data = np.array(data)
    cluster_centers = np.array(cluster_centers)
    if cluster_centers is not None:
        wss = 0
        total_points = 0
        for i in range(len(cluster_centers)):
            # Extract points belonging to the current cluster
            cluster_points = clustered_data[clustered_labels == i]
            total_points += len(cluster_points)
            for j in range(len(cluster_points)):
                squared_distance = np.sum((cluster_points[j] - cluster_centers[i]) ** 2, axis=0)
                wss += squared_distance
        if total_points > 0:
            weighted_wss = wss / total_points
        else:
            weighted_wss = np.nan
    else:
        weighted_wss = np.nan
    # Calculating BIC
    if model is not None:
        n_clusters = len(np.unique(clustered_labels))
        n_features = data.shape[1]
        n_samples = len(clustered_data)
        # Log likelihood approximation for BIC
        if cluster_centers is not None:
            distances = cdist(clustered_data, cluster_centers)
            min_distances = np.min(distances, axis=1)
            log_likelihood = -0.5 * np.sum(min_distances ** 2, axis=0)
        else:
            log_likelihood = np.nan
        # BIC calculation
        n_params = n_clusters * n_features  # Approximate number of parameters
        bic = -2 * log_likelihood + n_params * np.log(n_samples)
    else:
        bic = np.nan  # BIC requires a model

    return {
        "Clustering Name": name,
        "Noise Percentage": noise_percentage,
        "Weighted WSS": weighted_wss,
        "Silhouette Score": silhouette,
        "BIC": bic
    }


def convert_pdf_to_images(file_path, output_dir):
    """
    Converts each page of a PDF to an image and save it in the specified directory.
    Args:
    - file_path (str): Path to the PDF file.
    - output_dir (str): Directory where the output images should be saved.
    Returns:
    - None
    """
    doc = fitz.open(file_path)
    for i, page in enumerate(doc):
        pix = page.get_pixmap()
        output_path = f"{output_dir}/PowerBI_{i}.png"
        pix.save(output_path)
    print(f"All pages have been converted to images and saved in {output_dir}.")


def cramers_v(x, y):
    """To calculate Cramér's V for two categorical variables"""
    contingency_table = pd.crosstab(x, y)
    chi2 = stats.chi2_contingency(contingency_table)[0]
    n = contingency_table.sum().sum()
    phi2 = chi2 / n
    r, k = contingency_table.shape
    return np.sqrt(phi2 / min(r - 1, k - 1))


def drop_high_corr_features(cm, threshold, var_table):
    high_corr_pairs = []
    features_to_drop = []
    variance_dict = dict(zip(var_table['features'], var_table['variances']))
    for _i_ in range(len(cm.columns)):
        for j in range(_i_):
            if abs(cm.iloc[_i_, j]) > threshold:
                high_corr_pairs.append((cm.columns[_i_], cm.columns[j]))
                feature_i = cm.columns[_i_]
                feature_j = cm.columns[j]
                if variance_dict[feature_i] < variance_dict[feature_j]:
                    features_to_drop.append(feature_i)
                else:
                    features_to_drop.append(feature_j)
    return high_corr_pairs, features_to_drop


def iterative_correlation_dropper(current_data, cutoff, varframe, min_features=8):
    """
    Iteratively drops features with high correlation, keeping the column with lower variance.
    Parameters:
        current_data (pd.DataFrame): The dataframe to filter.
        cutoff (float): Correlation threshold for dropping columns.
        varframe (pd.DataFrame): Dataframe containing feature variances.
        min_features (int): Minimum number of features to retain.
    Returns:
        pd.DataFrame: The filtered dataframe with reduced correlation.
    """
    all_dropped_features = []
    all_corr_pairs = []
    while len(current_data.columns) > min_features:
        # Calculating the correlation matrix
        corr_matrix = current_data.corr(method='spearman')
        # Masking to ignore redundant correlations
        mask = np.triu(np.ones(corr_matrix.shape), k=0)
        corr_matrix = corr_matrix.where(mask == 0)
        # Finding pairs with correlation above cutoff
        corr_pairs = corr_matrix.stack()
        vs_corr_pairs = corr_pairs[abs(corr_pairs) > cutoff].sort_values(ascending=False, key=abs)
        if vs_corr_pairs.empty:
            print(f"No more pairs above an absolute correlation of {cutoff}. Stopping.")
            break
        f1, f2 = vs_corr_pairs.index[0]
        if f1 not in current_data.columns or f2 not in current_data.columns:
            continue
        var_f1 = varframe.loc[varframe['features'] == f1, 'variances'].values[0]
        var_f2 = varframe.loc[varframe['features'] == f2, 'variances'].values[0]
        feature_to_drop = f1 if var_f1 > var_f2 else f2
        current_data = current_data.drop(columns=[feature_to_drop])
        all_dropped_features.append(feature_to_drop)
        all_corr_pairs.append((f1, f2))
        print(f"Dropped feature: {feature_to_drop} (Correlation: {vs_corr_pairs.iloc[0]:.3f})")
        # Stopping condition to retain a minimum number of features
        if len(current_data.columns) <= min_features:
            print(f"Reached the minimum number of features: {min_features}. Stopping.")
            break
    print(f"Final feature count: {len(current_data.columns)}")
    return current_data


def missing_value_ratio(col):
    return (col.isnull().sum() / len(col)) * 100


# Creating shortened names for columns
short_labels = {
    'Gender': 'Gen',
    'Country': 'Continent',
    'self_employed': 'SelfEmp',
    'family_history': 'FamHist',
    'treatment': 'Treat',
    'remote_work': 'Remote',
    'benefits': 'Benefits',
    'care_options': 'CareOpt',
    'wellness_program': 'WellProg',
    'seek_help': 'SeekHelp',
    'anonymity': 'Anonymity',
    'leave': 'Leave',
    'mental_health_consequence': 'MenCons',
    'phys_health_consequence': 'PhysCons',
    'coworkers': 'Coworkers',
    'supervisor': 'Supervisor',
    'mental_health_interview': 'MentInt',
    'phys_health_interview': 'PhysInt',
    'mental_vs_physical': 'MentvsPhys',
    'obs_consequence': 'ObsCons'
}


def save_as_word_table(dataframe, file_name):
    doksi = Document()
    doksi.add_heading('Categorical Feature Summary', level=1)
    table = doksi.add_table(rows=1, cols=len(dataframe.columns))
    table.style = 'Table Grid'
    for idx, column in enumerate(dataframe.columns):
        table.cell(0, idx).text = column
    for _, row in dataframe.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    doksi.save(file_name)


def top_10_largest_ranges(szumma_df):
    summary = szumma_df.describe(include='all').transpose()
    summary = summary[['mean', 'min', 'max']]
    summary['range'] = summary['max'] - summary['min']
    largest_ranges = summary.sort_values(by='range', ascending=False).head(10)
    print("Features with the 10 Largest Ranges:")
    print(tabulate.tabulate(largest_ranges, headers=['Feature', 'Mean', 'Min', 'Max', 'Range'], tablefmt='pretty'))


def variance_testing(dframe, varthresh):
    selector = VarianceThreshold(threshold=varthresh)
    _ = selector.fit_transform(dframe)
    variances = selector.variances_
    variance_df = pd.DataFrame({'features': dframe.columns, 'variances': variances})
    variance_df = variance_df.sort_values(by='variances', ascending=False)
    selected_features = dframe.columns[selector.get_support(indices=True)]
    return variance_df, selected_features


def winsorization_outliers(df):
    out = []
    for n in df:
        q1 = np.percentile(df, 1)
        q3 = np.percentile(df, 99)
        if n > q3 or n < q1:
            out.append(n)
    print("Outliers:", out)
    return out

# file_path = ("D:/Tanulás/iu/Subjects/Machine Learning - Unsupervised Learning and Feature "
#              "Engineering/PowerBI/untitled.pdf")
# doc = fitz.open(file_path)
# for i, page in enumerate(doc):
#     pix = page.get_pixmap()  # render page to an image
#     pix.save(f"D:/Tanulás/iu/Subjects/Machine Learning - Unsupervised Learning and Feature "
#              "Engineering/PowerBI/scaleds_{i}.png")
